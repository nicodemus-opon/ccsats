import sqlite3
import datetime
import zlib
import zipfile
import hashlib
import shutil
import csv
import os
import random
from docx.shared import Inches
from flask import Flask, render_template, redirect, url_for, request, session, send_from_directory
from functools import wraps
import docx
from docx.api import Document
from werkzeug.utils import secure_filename
from main import *
UPLOAD_FOLDER = 'static/merge/'
ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'png', 'jpg', 'jpeg', 'docx'])
#from flask_mysqldb import MySQL
#from flaskext.mysql import MySQL

app = Flask(__name__)
app.secret_key = "nico"
#mysql = MySQL()
print("initializing database")
app.config['MYSQL_USER'] = 'sql9259727' #'root'
app.config['MYSQL_PASSWORD'] = '7D7EWBCl1r'#"'123'"
app.config['MYSQL_DB'] = 'sql9259727'#'ccsats'
app.config['MYSQL_PORT'] = 3306
app.config['MYSQL_DATABASE_HOST']='fsql9.freemysqlhosting.net'
app.config['MYSQL_CURSORCLASS'] = 'DictCursor'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#mysql = MySQL(app)
#con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
print("database initialized")
#mysql.init_app(app)
app.debug = True
import pymysql as mysql
import pymysql
def is_logged_in(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'logged_in' in session:
            return f(*args, **kwargs)
        else:
            return redirect(url_for('index'))

    return wrap

@app.route('/')
def index():
    return (render_template('index.html'))

def namer(nm=""):
    #replace name in fist page
    replaceText = {"NAME" : nm}
    # template directory
    templateDocx = zipfile.ZipFile("static/merge/template.docx")
    # final document
    newDocx = zipfile.ZipFile("static/merge/nicopon.docx", "a") 
    # temporary diectory
    with open(templateDocx.extract("word/document.xml", "static/merge/")) as tempXmlFile:
        tempXmlStr = tempXmlFile.read()
    #print(tempXmlStr)
    print(type(tempXmlStr))
    for key in replaceText.keys():
        tempXmlStr = tempXmlStr.replace(str(key), str(replaceText.get(key)))
    with open("static/merge/temp.xml", "w+") as tempXmlFile:
        tempXmlFile.write(tempXmlStr)

    for file in templateDocx.filelist:
        if not file.filename == "word/document.xml":
            newDocx.writestr(file.filename, templateDocx.read(file))

    newDocx.write("static/merge/temp.xml", "word/document.xml")

    templateDocx.close()
    newDocx.close()
    ##################################################################
    WORKING_DIR = os.getcwd()
    TEMP_DOCX = os.path.join(WORKING_DIR, "static/merge/nicopon.docx")
    TEMP_ZIP = os.path.join(WORKING_DIR, "static/merge/nicopon.zip")
    TEMP_FOLDER = os.path.join(WORKING_DIR, "static/merge/nicopon")
    if os.path.exists(TEMP_ZIP):
        os.remove(TEMP_ZIP)
    if os.path.exists(TEMP_FOLDER):
        shutil.rmtree(TEMP_FOLDER)
    os.rename(TEMP_DOCX, TEMP_ZIP)
    with zipfile.ZipFile(TEMP_ZIP, 'r') as z:
        z.extractall(TEMP_FOLDER)
    x=1
    while x<5:
        names="header"+str(x)+".xml"
        print(names)
        header_xml = os.path.join(TEMP_FOLDER, "word", names)
        try:
            header_xml = os.path.join(TEMP_FOLDER, "word", names)
            break
        except Exception as e:
            x+=1
            print(e)
    xmlstring = open(header_xml, 'r', encoding='utf-8').read()
    xmlstring = xmlstring.replace("NAME", nm)
    with open(header_xml, "wb") as f:
        f.write(xmlstring.encode("UTF-8"))
    os.remove(TEMP_ZIP)
    shutil.make_archive(TEMP_ZIP.replace(".zip", ""), 'zip', TEMP_FOLDER)
    os.rename(TEMP_ZIP, TEMP_DOCX)
    shutil.rmtree(TEMP_FOLDER)


def read_table(docv):
    document = Document(docv)
    table = document.tables[1]
    table_info=document.tables[0]
    data = []
    info=[]
    keys = None
    for i, row in enumerate(table_info.rows):
        text = [cell.text for cell in row.cells]
        info.append(text)
        
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        tr=[cell.text for cell in row.cells]
        hu=len(tr)
        if hu==4:
            keys=("number","title","comments","no")
        elif hu==3:
            keys=("number","title","comments")
        elif hu==2:
            keys=("title","comments")    
        row_data = dict(zip(keys, text))
        data.append(row_data)
    info=[info[0][1],info[0][3],info[1][1],info[1][3]]
    namer(info[0])
    session["inf"]=info
    #info candidate,excercise,assesor,date  would be good if stored in sessions
    return data

def merger(docs=[]):
    list_of_data=[]
    list_of_names=[]
    for z in docs:
        name=z.split(".")
        na=name[0]
        list_of_names.append(na)
    list_of_names=session["docies"]
    for x in docs:
        list_of_data.append(read_table(x))
    documentx = Document("static/merge/nicopon.docx")
    documentx.add_page_break()
    y=0
    x=0
    op=list_of_data
    list_of_comments=[]
    k=0
    m=0
    for x in range(len(op)):
        templist=[]
        for y in op[x]:
            kk=y['comments']
            kk=str(kk).replace("+","")
            kk=str(kk).replace("-","")
            kk=str(kk).replace("-/+","")
            kk=str(kk).replace("+/-","")
            templist.append(kk)
        list_of_comments.append(templist)
        
    list_of_titles=[]
    for x in range(len(op)):
        templist=[]
        print(x)
        for y in op[x]:
            print(y['title'])
            templist.append(y['title'])
        list_of_titles.append(templist)
    post=0
    print(list_of_titles)
    for x in range(len(list_of_titles)):
        for y in range(len(list_of_titles[x])):
            if x==0:
                if list_of_titles[x][y] in list_of_titles[1]:
                    nn=list_of_titles[x][y]
                    ind=list_of_titles[x+1].index(nn)
                    documentx.add_heading(list_of_titles[x][y], level=1)
                    i=0
                    while i<len(list_of_titles):
                        head=list_of_names[i]+'\n'
                        p = documentx.add_paragraph('\n')
                        p.add_run(head).bold = True
                        if i==0:
                            p.add_run(list_of_comments[x][y])
                        else:
                            p.add_run(list_of_comments[x+1][ind])
                        i+=1
                    documentx.add_page_break()
                else:
                    documentx.add_heading(list_of_titles[x][y], level=1)
                    head=list_of_names[x]+'\n'
                    p = documentx.add_paragraph('\n')
                    p.add_run(head).bold = True
                    p.add_run(list_of_comments[x][y])
                    documentx.add_page_break()
            else:
                if list_of_titles[x][y] in list_of_titles[0]:
                    pass
                else:
                    documentx.add_heading(list_of_titles[x][y], level=1)
                    head=list_of_names[x]+'\n'
                    p = documentx.add_paragraph('\n')
                    p.add_run(head).bold = True
                    p.add_run(list_of_comments[x][y])
                    documentx.add_page_break()
    p = documentx.add_paragraph()
    r = p.add_run()
    r.add_picture('static/end.PNG', width=Inches(7.09), height=Inches(8.76))
    documentx.save('static/merge/report.docx')


@app.route('/merger', methods=['POST','GET'])
def test():
    if request.method == "POST":
        docies=request.form['docies']
        namer=request.form['name']
        date=request.form['date']
        assesor=request.form['assesor']
        docies=docies[:-1]
        print(docies)
        session["docies"]=docies.split(".")
        filer = request.files.getlist('template')
        print(str(filer))
        print("huy")
        for file in filer:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filename="template.docx"
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        filex = request.files.getlist('file')
        list_of_file=[]
        print(filex)
        for file in filex:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                xc=os.path.join(app.config['UPLOAD_FOLDER'], filename)
                list_of_file.append(xc)
        print("got all input")
        merger(list_of_file)
        cool_name=str(session["inf"][0])+" report.docx"
        return send_from_directory(directory='static/merge', filename="report.docx", as_attachment=True,attachment_filename=cool_name)
    folder = 'static/merge'
    for the_file in os.listdir(folder):
        file_path = os.path.join(folder, the_file)
        try:
            if os.path.isfile(file_path) and the_file!="template.docx" and the_file!="main.txt" and the_file.lower()!="end.png":
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(e)
    return (render_template('merger.html'))

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/', methods=['POST','GET'])
def process_login():
    if request.method == "POST":
        erro=""
        email = request.form['email']
        password = request.form['password']
        print(email)
        print(password)
        #con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        con= mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        #con=mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute("SELECT * FROM users")
        with con:
            rows = cur.fetchall()
            print("===================")
            print(rows[0])
            print(rows[1])
            print("===================")
            for row in rows:
                print("===")
                print(row['usrname'])
                print(row['passwrd'])
                print("===")
                if row['usrname']==email and row['passwrd']==password:
                    session["username"]=row['usrname']
                    session['logged_in']=True
                    session["erro"]=erro
                    session.pop("erro")
                    return redirect(url_for('dashboard'))
                else:
                    print("invalid creds")
        erro="Invalid Credentials !"
        session["erro"]=erro
        return render_template('index.html')
    return render_template('index.html')

@app.route('/mysurveys')
def mysurveys():
    con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
    cur = con.cursor()
    cur.execute("SELECT nme FROM survey where usrname='"+str(session["username"])+"'")
    list_of_surveysx=[]
    list_of_respos=[]
    list_of_urls=[]
    list_of_dels=[]
    with con:
        rows = cur.fetchall()
        for row in rows:
            list_of_surveysx.append(str(row['nme']))
            urls="/survey/"+str(row['nme'])
            dels="/n/"+str(row['nme'])
            list_of_dels.append(dels)
            list_of_urls.append(urls)
    session["snames"]=list_of_surveysx
    session["lens"]=len(list_of_surveysx)
    session["sview"]=list_of_urls
    session["sdelete"]=list_of_dels
    for x in list_of_surveysx:
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute("SELECT * FROM "+str(x).lower())
        with con:
            rows = cur.fetchall()
            vg=0
            for row in rows:
                vg+=1
                #list_of_values.append(list(row.values()))
        list_of_respos.append(vg)
    session["sresponses"]=list_of_respos
    print(session["snames"])
    print(session["sview"])
    print(session["sdelete"])
    print(session["sresponses"])
    return (render_template('mysurveys.html'))

@app.route('/design', methods=['GET', 'POST'])
@is_logged_in
def design():
    if request.method == 'POST':
        print("gotten some stuff")
        return("got sometin")
    return (render_template('design.html'))

@app.route('/dashboard', methods=['GET', 'POST'])
@is_logged_in
def dashboard():
    if request.method == 'POST':
        session["surveyname"]=request.json
        return (redirect(url_for("table")))
    else:
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        ds="SELECT * FROM notifications where usrname='"+session["username"]+"';"
        cur.execute(ds)
        list_of_vals=[]
        with con:
            rows = cur.fetchall()
            for row in rows:
                list_of_vals.append(row["alert"])
        session["alerts"]=list_of_vals
        session["lenalerts"]=len(list_of_vals)
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute("SELECT nme FROM survey where usrname='"+str(session["username"])+"'")
        list_of_surveys=[]
        list_of_cols=[]
        session["surveyname"]='notable'
        with con:
            rows = cur.fetchall()
            for row in rows:
                list_of_surveys.append(str(row['nme']))
                session["surveyname"]=row['nme']
        print(list_of_surveys)
        session["list_of_surveys"]=list_of_surveys
        session["number_of_surveys"]=len(list_of_surveys)
        return (render_template('dashboard.html'))


@app.route('/success')
def success():
    return (render_template('success.html'))

@app.route('/publish')
def successv():
    return (render_template('publish.html'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    erro = ""
    session["erro"]=erro
    session.pop("erro")
    if request.method == 'POST':
        print("get req")
        emailreg = request.form['emailreg']
        passwordreg = request.form['passwordreg']
        print("====register=====")
        print(emailreg)
        print(passwordreg)
        print("====+++++++=====")
        que='''select * from users;'''
        quert='''insert into users values("'''+emailreg+'''","'''+passwordreg+'''");'''
        print(quert)
        #cur = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)ion.cursor()
        #cur.execute(quert)
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute(quert)
        con.commit()
        print("registered")
        #rv = cur.fetchall()
        #print(rv)
        return (redirect(url_for('index')))
    return render_template('register.html')

@app.route('/parse_data', methods=['GET', 'POST'])
def parse_data():
    datax=None
    if request.method == "POST":
        print("got post")
        datax=request.json
        dat_array=datax.split("*")
        title=dat_array[0]
        header=dat_array[1]
        survey_name=dat_array[2]
        ht=dat_array[3]
        lent=dat_array[4]
        ht=ht.replace("  ","")
        ht=ht.replace('"',"|")
        ht=ht.replace("'","|")
        session["title"]=title
        session["header"]=header
        session["survey_name"]=survey_name
        session["ht"]=ht
        session["lent"]=lent
        session["urlo"]=str(request.host_url)+"survey/"+str(session["survey_name"])
        print(session["urlo"])
        print(session["title"])
        print(session["header"])
        print(session["survey_name"])
        print(session["ht"])
        print(session["lent"])
        length=int(session["lent"])
        init_string="create table "+ str(session["survey_name"]) +"("
        for x in range(length):
            num=x+1
            textpend="q"+str(num)+" text ,"
            init_string+=textpend
        init_string=init_string[:-1]
        init_string+=");"
        init_string=init_string.lower()
        print(init_string)
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute(init_string)
        con.commit()
        print("added table")
        po=str(session["survey_name"])
        po=po.lower()
        quert='''insert into survey values("'''+po+'''","'''+session["username"]+'''","''' + session["ht"]+'''","''' + session["title"]+'''","''' + session["header"]+'''");'''
        #print(quert)
        #cur = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)ion.cursor()
        #cur.execute(quert)
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        #quert=quert.lower()
        cur.execute(quert)
        con.commit()
        print("registered survey")
        redstring="/survey/"+session["survey_name"]

        redstring="/publish"
        return(render_template("publish.html"))

def download_csv(name):
    mystring=name
    hash_object = hashlib.md5(mystring.encode())
    #hash_object.hexdigest()
    filex="static/downloads/"+str(session["surveyname"])+".csv"
    xc=str(session["surveyname"])
    session["filex"]=filex
    with open(filex, "w",newline='') as f:
        writer = csv.writer(f)
        writer.writerows(session["mentions"])
    return(xc)


@app.route('/table', methods=['GET', 'POST'])
def table():
    if request.method == 'POST':
        print("got post")
        datax=session["surveyname"]
        print(datax)
        filenamex=download_csv(datax)
        print(filenamex)
        pathx="/download/"+filenamex
        return(redirect(pathx))
    session['dat']=[]
    session['cols']=[]
    con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
    cur = con.cursor()
    cur.execute("SELECT * FROM "+str(session["surveyname"]).lower())
    list_of_values=[]
    list_of_cols=[]
    with con:
        rows = cur.fetchall()
        for row in rows:
            list_of_values.append(list(row.values()))
    print("***********************")
    print(list_of_values)
    print("***********************")
    if list_of_values==[]:
        list_of_values=[['UH OH ;)'],['This Survey has not been filled']]
    for z in range(len(list_of_values[0])):
        e=z+1
        d="que-"+str(e)
        list_of_cols.append(d)
    
    session["mentions"]=list_of_values
    session["cols"]=list_of_cols
    session["mentionsx"]=len(list_of_values)
    session["mentionsxy"]=len(list_of_values[0])
    print(session["mentions"])
    return render_template("table.html")

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        completion = validate(username, password)
        if completion == False:
            app.logger.info("PASSWORD NOT MATCHED")
            error = 'Invalid Credentials. Please try again...'
        else:
            app.logger.info("PASSWORD MATCHED")
            session['logged_in'] = True
            session['username'] = username
            con = sqlite3.connect('static/login.db')
            with con:
                cur = con.cursor()
                cur.execute("SELECT * FROM users")
                rows = cur.fetchall()
                for row in rows:
                    dbUser = row[0]
                    dbsub = row[3]
                    if dbUser == username:
                        subscription = dbsub
            session['subscription'] = subscription
            return redirect(url_for('dashboard'))
    return render_template('login.html', error=error)


def validate(username, password):
    con = sqlite3.connect('static/login.db')
    completion = False
    with con:
        cur = con.cursor()
        cur.execute("SELECT * FROM users")
        rows = cur.fetchall()
        for row in rows:
            dbUser = row[0]
            dbPass = row[1]
            if dbUser == username:
                completion = check_password(dbPass, password)
    return completion


def check_password(hashed_password, user_password):
    return hashed_password == user_password


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))


@app.route('/uh-oh')
def uhoh():
    return render_template("uhoh.html")


@app.route('/merge', methods=['GET', 'POST'])
def merge():
    if request.method == 'POST':
        '''if 'file' not in request.files:
            print('No file part')
            return redirect(request.url)
        
        if file.filename == '':
            print('No selected file')
            return redirect(request.url)'''
        filer = request.files.getlist('file')
        list_of_file=[]
        for file in filer:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                #filename="meeting.docx"
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                xc=os.path.join(app.config['UPLOAD_FOLDER'], filename)
                list_of_file.append(xc)
        print(list_of_file)
        merger(list_of_file)
        return send_from_directory(directory='static/merge', filename="report.docx", as_attachment=True)
    return render_template("merge.html")

def compress(file_names):
    print("File Paths:")
    print(file_names)

    path = "/static/downloads/"
    # Select the compression mode ZIP_DEFLATED for compression
    # or zipfile.ZIP_STORED to just store the file
    compression = zipfile.ZIP_DEFLATED

    # create the zip file first parameter path/name, second mode
    zf = zipfile.ZipFile("/static/downloads/RAWs.zip", mode="w")
    try:
        for file_name in file_names:
            # Add file to the zip file
            # first parameter file to zip, second filename in zip
            zf.write(path + file_name, file_name, compress_type=compression)

    except FileNotFoundError:
        print("An error occurred")
    finally:
        # Don't forget to close the file!
        zf.close()


@app.route('/downloadall')
def downloadall():
    list_of_files=[]
    for datax in session["list_of_surveys"]:
        print("downloading...")
        print(datax)
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute("SELECT * FROM "+str(datax))
        list_of_values=[]
        list_of_cols=[]
        with con:
            rows = cur.fetchall()
            for row in rows:
                list_of_values.append(list(row.values()))
        #filenamex=download_csv(datax)
        filex="static/downloads/"+str(datax)+".csv"
        xc=str(datax)
        session["filex"]=filex
        with open(filex, "w",newline='') as f:
            writer = csv.writer(f)
            writer.writerows(list_of_values)
        fij=str(datax)+".csv"
        list_of_files.append(fij)
        #print(filenamex)
        #pathx="/download/"+filenamex
        #print(pathx)
        #return(redirect(pathx))
    #compress(list_of_files)
    fulldir="static/downloads/"+str(session["username"])+".zip"
    filename=str(session["username"])+".zip"
    with zipfile.ZipFile(fulldir, 'w') as zipMe:        
        for file in list_of_files:
            file="static/downloads/"+file
            print(file)
            zipMe.write(file, compress_type=zipfile.ZIP_DEFLATED)
        
    #return redirect(url_for("dashboard"))
    return send_from_directory(directory='static/downloads', filename=filename, as_attachment=True)


@app.route('/notification', methods=['GET', 'POST'])
def notification():
    if request.method=="POST":
        print("removing notifications")
        qe="delete from notifications"
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute(qe)
        con.commit()
        return redirect(url_for("notification"))
    else:
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        ds="SELECT * FROM notifications where usrname='"+session["username"]+"';"
        cur.execute(ds)
        list_of_vals=[]
        with con:
            rows = cur.fetchall()
            for row in rows:
                list_of_vals.append(row["alert"])
        session["alerts"]=list_of_vals
        session["lenalerts"]=len(list_of_vals)
        return render_template("notification.html")


@app.route('/download/<string:filename>', methods=['GET', 'POST'])
def download(filename):
    print("got download")
    filename=filename+".csv"
    print(filename)
    return send_from_directory(directory='static/downloads', filename=filename, as_attachment=True)


@app.route("/survey/<string:name>")
def surve(name):
    session['surveyname']=name
    select_query="select * from survey where nme='"+str(name)+"' ;"
    con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
    cur = con.cursor()
    cur.execute(select_query)
    with con:
        rows = cur.fetchall()
        for row in rows:
            session['nme']=row['nme']
            session['tito']=str(row['title'])
            xc=row['html']
            xc=xc.replace("|",'"')
            xc=xc.replace("wrap",'')
            xc=xc.replace("<ul",'<li')
            xc=xc.replace("</ul",'</li')
            xc=xc.replace('<i class="now-ui-icons ui-1_simple-remove text-dark"></i>',"")
            session['heda']=row['header']
            session['htmlo']=xc
    return redirect(url_for("survey"))

@app.route("/t/<string:name>")
def t(name):
    session["surveyname"]=str(name)
    return redirect(url_for("table"))


@app.route("/n/<string:name>")
def n(name):
    qe="delete from survey where nme='"+name+"'"
    qu="drop table "+name
    con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
    cur = con.cursor()
    cur.execute(qe)
    con.commit()
    con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
    cur = con.cursor()
    cur.execute(qu)
    con.commit()
    print("deleted")
    return redirect(url_for("mysurveys"))


@app.route("/survey", methods=['GET', 'POST'])
def survey():
    if request.method == 'POST':
        datax=request.json
        print(datax)
        #print(datax[1])
        timestr=str(datetime.datetime.now().strftime("%I:%M%p on %B %d, %Y"))
        fullstr="survey "+str(session['surveyname'])+" was filled at "+timestr
        qstring=""
        #######
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute("select usrname from survey where nme='"+session["surveyname"]+"'")
        usri=""
        with con:
            rows = cur.fetchall()
            for row in rows:
                usri+=str(row["usrname"])
        #######
        execstr="insert into notifications values('"+usri+"','"+ fullstr+"');"
        for x in datax:
            qstring+="'"+x+"',"
        qstring=qstring[:-1]
        print(qstring)
        query="insert into "+str(session['surveyname'])+" values("+qstring+");"
        print(query)
        con = mysql.connect("db4free.net","ccsats","Ccsats11060!","ccsats",cursorclass=pymysql.cursors.DictCursor)
        cur = con.cursor()
        cur.execute(query)
        cur.execute(execstr)
        con.commit()
        print("complete")
    return render_template("survey.html")
'''
@app.route('/<string:page_name>/')
def render_static(page_name):
    return render_template('%s.html' % page_name)'''

@app.route("/dashboard/<string:name>")
def handle_dash(name):
    if name=="table":
        return redirect(url_for("table"))


@app.errorhandler(Exception)
def handle_error(e):
    print(e)
    e=str(e)
    session["error"]=e
    return render_template("uhoh.html")
    #code = 500
    #if isinstance(e, HTTPException):
    #    code = e.code
    #return jsonify(error=str(e)), code
    
if __name__ == '__main__':
    app.run(host='0.0.0.0', port= 8090)
